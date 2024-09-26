import os
from docx import Document
import re
import uuid

# Function to convert markdown text to docx format and return the file path with a random UUID
def markdown_to_docx(file_path):
    # Generate a random UUID
    random_uuid = str(uuid.uuid4())

    # Define the folder where you want to save the file
    upload_folder = os.path.join(os.path.dirname(file_path), "uploadable files")
    
    # Create the folder if it doesn't exist
    if not os.path.exists(upload_folder):
        os.makedirs(upload_folder)

    # Create a new .docx file name with the random UUID inside the uploadable files folder
    docx_file = os.path.join(upload_folder, f"{random_uuid}.docx")

    # Read the markdown (.txt) file content
    with open(file_path, "r", encoding="utf-8") as file:
        md_content = file.readlines()

    # Create a new Document object for .docx
    doc = Document()

    # Patterns to detect markdown syntax
    header_pattern = re.compile(r'^(#+)\s+(.*)')
    list_pattern = re.compile(r'^[-*]\s+(.*)')
    bold_pattern = re.compile(r'\*\*(.*?)\*\*')
    italic_pattern = re.compile(r'\*(.*?)\*')

    for line in md_content:
        # Remove any leading/trailing whitespace
        line = line.strip()

        # Detect headers based on the number of '#' symbols
        header_match = header_pattern.match(line)
        if header_match:
            header_level = len(header_match.group(1))  # Number of '#' to determine header level
            header_text = header_match.group(2)
            doc.add_heading(header_text, level=header_level)
            continue

        # Detect list items (unordered)
        list_match = list_pattern.match(line)
        if list_match:
            list_item = list_match.group(1)
            doc.add_paragraph(list_item, style='ListBullet')
            continue

        # Apply bold and italic formatting
        bold_text = bold_pattern.sub(r'\1', line)  # Replace bold markdown with corresponding text
        italic_text = italic_pattern.sub(r'\1', bold_text)  # Replace italic markdown with corresponding text
        doc.add_paragraph(italic_text)

    # Save the document with the random UUID name
    doc.save(docx_file)
    print(f"Document saved as {docx_file}")

    # Return the generated file path
    return docx_file
